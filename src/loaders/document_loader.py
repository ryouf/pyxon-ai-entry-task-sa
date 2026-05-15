import re
from pathlib import Path
from src.processing.arabic_processor import clean, get_stats
import fitz
import pytesseract
import cv2
import numpy as np
from PIL import Image
import io
from src.processing.arabic_processor import clean, get_stats, clean_ocr
  
def load_pdf(file_path):
    
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

    pages = []
    doc = fitz.open(file_path)
    meta = {
        "title": doc.metadata.get("title", "") or Path(file_path).stem,
        "author": doc.metadata.get("author", ""),
        "num_pages": len(doc),
    }

    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=400)
        img_array = np.frombuffer(pix.samples, dtype=np.uint8)
        img_cv = img_array.reshape(pix.height, pix.width, pix.n)

        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        gray = cv2.medianBlur(gray, 3)

        img_pil = Image.fromarray(gray)

        custom_config = r'--oem 3 --psm 6 -l ara+eng'
        text = pytesseract.image_to_string(img_pil, config=custom_config)
        text = clean_ocr(text)
        text = clean(text)

        if text:
            pages.append({"page": i + 1, "text": text})

    doc.close()
    return meta, pages

def load_docx(file_path):
    from docx import Document
    doc = Document(file_path)
    meta = {
        "title": doc.core_properties.title or Path(file_path).stem,
        "author": doc.core_properties.author or "",
        "num_pages": len(doc.paragraphs),
    }
    pages = []
    current = []
    section = 1
    for para in doc.paragraphs:
        text = clean(para.text)
        if not text:
            continue
        if para.style and "Heading" in para.style.name and current:
            pages.append({"page": section, "text": "\n".join(current)})
            current = []
            section += 1
        current.append(text)
    if current:
        pages.append({"page": section, "text": "\n".join(current)})
    return meta, pages

def load_txt(file_path):
    encodings = ["utf-8", "utf-8-sig", "windows-1256", "iso-8859-6", "cp1256"]
    text = ""
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                text = f.read()
            break
        except:
            continue
    text = clean(text)
    meta = {"title": Path(file_path).stem, "author": "", "num_pages": 1}
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    pages = [{"page": i + 1, "text": p} for i, p in enumerate(paragraphs)]
    return meta, pages

def load_document(file_path):
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        meta, pages = load_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        meta, pages = load_docx(file_path)
    elif suffix == ".txt":
        meta, pages = load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    full_text = "\n\n".join(p["text"] for p in pages)
    stats = get_stats(full_text)

    return {
        "meta": meta,
        "pages": pages,
        "full_text": full_text,
        "language": "arabic" if stats["is_arabic"] else "english",
        "has_diacritics": stats["has_diacritics"],
        "total_chars": stats["total_chars"],
        "file_type": suffix.lstrip("."),
        "file_path": str(file_path),
    }